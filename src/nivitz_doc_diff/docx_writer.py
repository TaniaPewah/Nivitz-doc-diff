"""Build output .docx with highlight/strikethrough formatting applied to diff regions."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from lxml import etree

from .diff_engine import DiffOp, DocDiffOp


def ensure_rpr(run):
    """Ensure the run has an rPr (run properties) element, creating one if needed."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
    return rpr


def apply_highlight(run, color_index=WD_COLOR_INDEX.YELLOW):
    """Apply background highlight color to a run."""
    rpr = ensure_rpr(run)
    hl = rpr.find(qn("w:highlight"))
    if hl is None:
        hl = etree.SubElement(rpr, qn("w:highlight"))
    hl.set(qn("w:val"), str(color_index))
    return run


def apply_strikethrough(run):
    """Apply strikethrough formatting to a run."""
    rpr = ensure_rpr(run)
    strike = rpr.find(qn("w:strike"))
    if strike is None:
        strike = etree.SubElement(rpr, qn("w:strike"))
    strike.set(qn("w:val"), "true")
    return run


def _add_formatted_text(paragraph, text: str, *, highlight: bool = False, strikethrough: bool = False):
    """Add a run to a paragraph with optional highlighting and/or strikethrough."""
    if not text:
        return None
    run = paragraph.add_run(text)
    if highlight:
        apply_highlight(run)
    if strikethrough:
        apply_strikethrough(run)
    return run


def _render_word_ops(paragraph, word_ops: list[DiffOp]) -> None:
    """Render a list of word-level diff ops into a paragraph."""
    for op in word_ops:
        if op.kind == "equal":
            _add_formatted_text(paragraph, op.new_text)
        elif op.kind == "insert":
            _add_formatted_text(paragraph, op.new_text, highlight=True)
        elif op.kind == "delete":
            _add_formatted_text(paragraph, op.old_text, strikethrough=True)
        elif op.kind == "replace":
            # Show deleted (strikethrough) then inserted (highlighted)
            _add_formatted_text(paragraph, op.old_text, strikethrough=True)
            _add_formatted_text(paragraph, op.new_text, highlight=True)


def build_diff_document(diff_results, output_path: str | Path) -> None:
    """Build a .docx file from diff results.

    Accepts either:
      - list[DocDiffOp]  (v0.2+ API, paragraph-level alignment)
      - list[list[DiffOp]]  (v0.1 legacy API, index-based alignment)

    The legacy path is preserved for backwards compatibility with older tests.
    """
    doc = Document()

    # Detect which API we're dealing with
    is_new_api = diff_results and isinstance(diff_results[0], DocDiffOp)
    is_legacy_api = diff_results and isinstance(diff_results[0], list)

    if is_new_api:
        for doc_op in diff_results:
            para = doc.add_paragraph()
            if doc_op.kind == "equal":
                new_p = doc_op.new_paragraph
                _add_formatted_text(para, new_p.full_text if new_p else "")
            elif doc_op.kind == "insert":
                new_p = doc_op.new_paragraph
                _add_formatted_text(para, new_p.full_text if new_p else "", highlight=True)
            elif doc_op.kind == "delete":
                old_p = doc_op.old_paragraph
                _add_formatted_text(para, old_p.full_text if old_p else "", strikethrough=True)
            elif doc_op.kind == "replace":
                _render_word_ops(para, doc_op.word_ops)

            # Ensure paragraph has at least one run
            if not para.runs:
                _add_formatted_text(para, "")

    elif is_legacy_api:
        # v0.1 API: list[list[DiffOp]] — each sublist is one paragraph's word ops
        for para_ops in diff_results:
            para = doc.add_paragraph()
            _render_word_ops(para, para_ops)
            if not para.runs:
                _add_formatted_text(para, "")

    else:
        # Empty diff: empty document
        pass

    doc.save(str(output_path))
