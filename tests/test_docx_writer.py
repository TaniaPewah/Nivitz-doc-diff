"""Tests for docx_writer.py."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from nivitz_doc_diff.diff_engine import DiffOp, DocDiffOp
from nivitz_doc_diff.docx_reader import ParagraphData, RunData, read_docx
from nivitz_doc_diff.docx_writer import apply_highlight, apply_strikethrough, build_diff_document


NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _para(text: str) -> ParagraphData:
    return ParagraphData(index=0, runs=[RunData(text=text, xml_element=None)])


def _collect_runs(doc_path: Path) -> list[tuple[str, list[str]]]:
    """Return [(text, [attr,...]), ...] for every run in the doc."""
    doc = Document(str(doc_path))
    result = []
    for para in doc.paragraphs:
        for run in para.runs:
            attrs = []
            rpr = run._element.find("{%s}rPr" % NS)
            if rpr is not None:
                if rpr.find("{%s}highlight" % NS) is not None:
                    attrs.append("HIGHLIGHT")
                strike = rpr.find("{%s}strike" % NS)
                if strike is not None and strike.get("{%s}val" % NS) == "true":
                    attrs.append("STRIKE")
            result.append((run.text, attrs))
    return result


# ---------- v0.2+ DocDiffOp API ----------


def test_build_diff_equal_paragraph(tmp_path: Path) -> None:
    ops = [DocDiffOp(
        kind="equal",
        old_paragraph=_para("Hello world"),
        new_paragraph=_para("Hello world"),
        word_ops=[DiffOp(kind="equal", old_text="Hello world", new_text="Hello world")],
    )]
    output = tmp_path / "diff.docx"
    build_diff_document(ops, output)

    result = read_docx(output)
    assert result.paragraphs[0].full_text == "Hello world"


def test_build_diff_insert_paragraph(tmp_path: Path) -> None:
    ops = [DocDiffOp(kind="insert", new_paragraph=_para("Brand new paragraph"))]
    output = tmp_path / "diff.docx"
    build_diff_document(ops, output)

    runs = _collect_runs(output)
    assert runs == [("Brand new paragraph", ["HIGHLIGHT"])]


def test_build_diff_delete_paragraph(tmp_path: Path) -> None:
    ops = [DocDiffOp(kind="delete", old_paragraph=_para("Removed paragraph"))]
    output = tmp_path / "diff.docx"
    build_diff_document(ops, output)

    runs = _collect_runs(output)
    assert runs == [("Removed paragraph", ["STRIKE"])]


def test_build_diff_replace_paragraph_word_level(tmp_path: Path) -> None:
    ops = [DocDiffOp(
        kind="replace",
        old_paragraph=_para("Hello old world"),
        new_paragraph=_para("Hello new world"),
        word_ops=[
            DiffOp(kind="equal", old_text="Hello ", new_text="Hello "),
            DiffOp(kind="replace", old_text="old", new_text="new"),
            DiffOp(kind="equal", old_text=" world", new_text=" world"),
        ],
    )]
    output = tmp_path / "diff.docx"
    build_diff_document(ops, output)

    runs = _collect_runs(output)
    # Must contain both struck "old" and highlighted "new"
    texts_by_attr = {}
    for text, attrs in runs:
        key = ",".join(attrs) if attrs else "plain"
        texts_by_attr.setdefault(key, []).append(text)
    assert any("old" in t for t in texts_by_attr.get("STRIKE", []))
    assert any("new" in t for t in texts_by_attr.get("HIGHLIGHT", []))


def test_build_diff_title_insertion_preserves_body(tmp_path: Path) -> None:
    """End-to-end: title inserted, body paragraph unchanged — body must be plain."""
    from nivitz_doc_diff.diff_engine import compute_diff

    old = [_para("Body text unchanged.")]
    new = [
        ParagraphData(index=0, runs=[RunData(text="Title", xml_element=None)]),
        ParagraphData(index=1, runs=[RunData(text="Body text unchanged.", xml_element=None)]),
    ]
    doc_ops = compute_diff(old, new)
    output = tmp_path / "diff.docx"
    build_diff_document(doc_ops, output)

    runs = _collect_runs(output)
    # Title should be highlighted (insert)
    title_runs = [(t, a) for t, a in runs if "Title" in t]
    assert any("HIGHLIGHT" in a for _, a in title_runs)
    # Body must be plain (unchanged)
    body_runs = [(t, a) for t, a in runs if "Body" in t]
    assert all(a == [] for _, a in body_runs), f"Body should be plain, got {body_runs}"


# ---------- v0.1 legacy list[list[DiffOp]] API still works ----------


def test_build_diff_legacy_list_of_diffops(tmp_path: Path) -> None:
    legacy = [[
        DiffOp(kind="equal", old_text="Hello ", new_text="Hello "),
        DiffOp(kind="insert", old_text="", new_text="beautiful "),
        DiffOp(kind="equal", old_text="world", new_text="world"),
    ]]
    output = tmp_path / "diff.docx"
    build_diff_document(legacy, output)

    runs = _collect_runs(output)
    highlighted = [t for t, attrs in runs if "HIGHLIGHT" in attrs]
    assert any("beautiful" in t for t in highlighted)
